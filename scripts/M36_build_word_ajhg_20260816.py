#!/usr/bin/env python3
# =============================================================================
# M36_build_word_ajhg_20260816.py — Assemble the AJHG-format submission Word
# document from the verified manuscript (manuscript.md), references, and the
# English-label figures.
#   Output: docs/manuscript/AJHG_submission_Qiushuo_Geng_20260816.docx
# Formatting per AJHG author guidelines (2026-08-16 revision):
#   - Author-edited title page format (title, author with superscript markers,
#     affiliations, #/* legend) rendered from the manuscript title block
#   - Times New Roman, 12 pt, double-spaced main text
#   - Numeric superscript citations ([@key] -> superscript number)
#   - Abstract <=200 words
#   - IMRaD sections, then Data and Code Availability, Web Resources,
#     Declaration of Interests, Funding
#   - References in NLM style (numbered by first citation order)
#   - Embedded figures (5 + S1), each with its full legend beneath (AJHG style)
#   - Three-line tables: Table 1 (data sources + yield summary, merged),
#     Table 2 (candidates), and supplemental Tables S1-S4 in Supplemental Items
# =============================================================================
import json, os, re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Repo root is self-locating so the script runs from any checkout.
BASE = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))
DOCS = f"{BASE}/docs/manuscript"
FIGD = f"{BASE}/results/figures"
OUTP = f"{DOCS}/AJHG_submission_Qiushuo_Geng_20260816.docx"

REFS = json.load(open(f"{DOCS}/refs.json"))

FIG_FILES = {
    1: f"{FIGD}/20260816_Fig1_design_genome.png",
    2: f"{FIGD}/20260816_Fig2_yield_funnel.png",
    3: f"{FIGD}/20260816_Fig3_calibration.png",
    4: f"{FIGD}/20260816_Fig4_regional.png",
    5: f"{FIGD}/20260816_Fig5_candidates.png",
}
S1_FILE = f"{FIGD}/20260816_FigS1_susie.png"

# combined inline tokenizer: **bold**, *italic*, [@key], <sup>…</sup>
INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[@[a-z0-9]+\]|<sup>[^<]*</sup>)")

def set_style_base(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

def _set_rfonts_tnr(rfonts):
    """Force every rFonts element to Times New Roman and drop theme-font attributes."""
    if rfonts is None:
        return
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), "Times New Roman")
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        a = qn(attr)
        if rfonts.get(a) is not None:
            del rfonts.attrib[a]

def force_times_new_roman(doc):
    """Post-build pass: make every text element render in Times New Roman.
    Word resolves fonts from (1) docDefaults, (2) the run's style, (3) the run's
    own rPr. python-docx's font.name only sets w:ascii/w:hAnsi, leaving the theme
    attributes (e.g. asciiTheme='minorHAnsi' = Calibri) in place, which Word
    prefers. This clears theme attributes everywhere and sets all four ranges."""
    # 1) docDefaults -> rPrDefault -> rFonts
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults"); styles_el.insert(0, dd)
    rpd = dd.find(qn("w:rPrDefault"))
    if rpd is None:
        rpd = OxmlElement("w:rPrDefault"); dd.append(rpd)
    rpr = rpd.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr"); rpd.append(rpr)
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    _set_rfonts_tnr(rf)

    # 2) every style (python-docx handles rPr insertion order for styles)
    for st in doc.styles:
        try:
            st.font.name = "Times New Roman"
        except Exception:
            continue
        rpr = st.element.find(qn("w:rPr"))
        if rpr is not None:
            rf = rpr.find(qn("w:rFonts"))
            if rf is not None:
                _set_rfonts_tnr(rf)

    # 3) every run in every paragraph (body + tables, incl. nested)
    def fix_para(p):
        for r in p.runs:
            rpr = r._r.get_or_add_rPr()
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts"); rpr.append(rf)
            _set_rfonts_tnr(rf)
    def fix_cell(c):
        for p in c.paragraphs:
            fix_para(p)
        for nt in c.tables:
            for row2 in nt.rows:
                for c2 in row2.cells:
                    fix_cell(c2)
    for p in doc.paragraphs:
        fix_para(p)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                fix_cell(c)

def add_run(p, text, bold=False, italic=False, superscript=False, font_size=12):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(font_size)
    r.font.bold = bold
    r.font.italic = italic
    if superscript:
        r.font.superscript = True
    return r

def set_double(p, after=0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(after)

def add_inline(p, text, font_size=12, base_bold=False):
    """Add text to an existing paragraph, handling **bold**, *italic*, [@key]."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            add_run(p, text[pos:m.start()], bold=base_bold, font_size=font_size)
        tok = m.group(1)
        if tok.startswith("[@"):
            key = tok[2:-1]
            num = REFS.get(key, "?")
            if num == "?": print(f"  !! missing citation key: {key}")
            add_run(p, str(num), bold=base_bold, superscript=True, font_size=font_size)
        elif tok.startswith("**"):
            add_run(p, tok[2:-2], bold=True, font_size=font_size)
        elif tok.startswith("*"):
            add_run(p, tok[1:-1], italic=True, font_size=font_size)
        elif tok.startswith("<sup>"):
            add_run(p, tok[5:-6], bold=base_bold, superscript=True, font_size=font_size)
        pos = m.end()
    if pos < len(text):
        add_run(p, text[pos:], bold=base_bold, font_size=font_size)

def add_body_para(doc, text, first_indent=None):
    p = doc.add_paragraph(); set_double(p)
    if first_indent: p.paragraph_format.first_line_indent = Inches(first_indent)
    add_inline(p, text)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); set_double(p)
    if level == 1:
        add_run(p, text, bold=True, font_size=13)
    else:
        add_run(p, text, bold=True, font_size=12)
    p.paragraph_format.keep_with_next = True
    return p

def set_cell_borders(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
    def set_edge(name, sz, val="single"):
        el = borders.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}"); borders.append(el)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
    if top is not None: set_edge("top", top)
    if bottom is not None: set_edge("bottom", bottom)

def add_three_line_table(doc, header, rows, font_size=10):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for j, h in enumerate(header):
        cell = t.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, h, font_size=font_size, base_bold=True)
        set_cell_borders(cell, top=12, bottom=6)
    for i, row in enumerate(rows, start=1):
        last = (i == len(rows))
        is_sub = row[0].strip().startswith("**")
        prev_is_sub = (i > 1) and rows[i - 2][0].strip().startswith("**")
        for j, v in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, v, font_size=font_size, base_bold=is_sub)
            if is_sub:
                set_cell_borders(cell, top=12, bottom=6)
            elif prev_is_sub:
                set_cell_borders(cell, bottom=12)
            else:
                set_cell_borders(cell, bottom=(12 if last else 0))
    return t

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    set_double(p)
    add_inline(p, text, font_size=11, base_bold=True)

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    set_double(p)
    add_inline(p, text, font_size=10)

def add_figure_page(doc, label, img_path, legend_text):
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(img_path, width=Inches(6.3))
    cap = doc.add_paragraph(); set_double(cap)
    add_inline(cap, f"**{label}.** " + legend_text, font_size=11)

def render_title_block(doc, md_text):
    """Render the manuscript title block (text before the first '---' separator)
    in the exact format of the author-edited DOCX title page:
    title (bold), author name (regular weight) with superscript affiliation
    markers, 'Affiliations:' label, affiliation lines, then the # / * legend."""
    block = md_text.split("\n---\n", 1)[0]
    for ln in block.split("\n"):
        s = ln.strip()
        if not s or s.startswith(">"):
            continue
        if s.startswith("# "):
            p = doc.add_paragraph(); set_double(p)
            add_run(p, s[2:].strip(), bold=True, font_size=14)
        elif s == "Affiliations:":
            p = doc.add_paragraph(); set_double(p)
            add_run(p, "Affiliations: ", bold=True, font_size=10.5)
        else:
            p = doc.add_paragraph(); set_double(p)
            # author name is regular weight in the target format: drop ** markers
            s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
            add_inline(p, s)

def parse_refs():
    refs = {}
    for line in open(f"{DOCS}/references.md", encoding="utf-8"):
        m = re.match(r"^(\d+)\.\s+(.*)", line.strip())
        if m: refs[int(m.group(1))] = m.group(2).strip()
    return refs

def split_sections(md_text):
    """Split manuscript.md into named sections; return list of (section_name, lines).
    Section boundaries are lines starting with '## '."""
    lines = md_text.split("\n")
    sections = []
    cur_name = None; cur = []
    for line in lines:
        m = re.match(r"^##\s+(.*)", line)
        if m:
            if cur_name is not None:
                sections.append((cur_name, cur))
            cur_name = m.group(1).strip(); cur = []
        else:
            if cur_name is not None:
                cur.append(line)
    if cur_name is not None:
        sections.append((cur_name, cur))
    return sections

def parse_paragraph_blocks(lines):
    """Yield ('h3'|'table'|'p'|'list'|'blank', payload) from section body lines."""
    blocks = []
    in_table = False; table_rows = []
    for line in lines:
        raw = line.rstrip()
        s = raw.strip()
        if not s:
            if in_table:
                blocks.append(("table", table_rows)); table_rows = []; in_table = False
            continue
        if s in ("---", "***", "___"):
            continue
        if s.startswith("|"):
            if not in_table: in_table = True; table_rows = []
            if not re.match(r"^\|[\s\-:|]+\|$", s):
                table_rows.append([c.strip() for c in s.strip("|").split("|")])
            continue
        if in_table:
            blocks.append(("table", table_rows)); table_rows = []; in_table = False
        if s.startswith("###"):
            blocks.append(("h3", s[3:].strip()))
            continue
        if s.startswith("- ") or s.startswith("* "):
            blocks.append(("list", s[2:].strip()))
            continue
        blocks.append(("p", s))
    if in_table:
        blocks.append(("table", table_rows))
    return blocks

def strip_figure_legend(text):
    """From a '**Figure N. <Title>.** <rest>' paragraph, return (num, '<Title>. <rest>').
    The 'Figure N.' prefix is added by add_figure_page, so it is NOT included here."""
    m = re.match(r"^\*\*Figure (\d+)\.\s*(.*?)\*\*(.*)$", text)
    if not m:
        return None, text
    num = int(m.group(1))
    title = m.group(2).strip().rstrip(".")
    rest = m.group(3).strip()
    body = f"{title}."
    if rest: body += " " + rest
    return num, body

def main():
    doc = Document()
    set_style_base(doc)
    for section in doc.sections:
        section.top_margin = Inches(1); section.bottom_margin = Inches(1)
        section.left_margin = Inches(1); section.right_margin = Inches(1)

    md_text = open(f"{DOCS}/manuscript.md", encoding="utf-8").read()
    sections = split_sections(md_text)
    by_name = {name: lines for name, lines in sections}

    # ---------- Title block (author-edited format: title, author with superscript
    #            affiliation markers, affiliations, #/* legend) ----------
    render_title_block(doc, md_text)
    doc.add_paragraph()

    # ---------- Main body (Abstract, numbered sections, data availability, web/declaration/funding) ----------
    BODY_ORDER = ["Abstract", "1. Introduction", "2. Methods", "3. Results",
                  "4. Discussion", "Data and Code Availability", "Web Resources",
                  "Declaration of Interests", "Funding"]
    for name in BODY_ORDER:
        if name not in by_name:
            print(f"  !! section not found: {name}")
            continue
        if name == "Abstract":
            add_heading(doc, "Abstract", 1)
        elif name.startswith(("1.", "2.", "3.", "4.")):
            add_heading(doc, name, 1)
        else:
            add_heading(doc, name, 1)
        for kind, payload in parse_paragraph_blocks(by_name[name]):
            if kind == "h3":
                add_heading(doc, payload, 2)
            elif kind == "p":
                add_body_para(doc, payload)
            elif kind == "list":
                p = doc.add_paragraph(); set_double(p)
                p.paragraph_format.left_indent = Inches(0.35)
                add_inline(p, payload)
            elif kind == "table":
                print("  !! table inside body section (unexpected):", payload[0][:4])

    # ---------- References ----------
    doc.add_page_break()
    add_heading(doc, "References", 1)
    refs_list = parse_refs()
    for num in sorted(refs_list):
        p = doc.add_paragraph(); set_double(p)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        add_run(p, f"{num}. "); add_run(p, refs_list[num])

    # ---------- Tables (from 'Tables' section) ----------
    if "Tables" in by_name:
        doc.add_page_break()
        add_heading(doc, "Tables", 1)
        for kind, payload in parse_paragraph_blocks(by_name["Tables"]):
            if kind == "p":
                add_caption(doc, payload)
            elif kind == "table":
                add_three_line_table(doc, payload[0], payload[1:])
            elif kind == "list":
                add_note(doc, payload)

    # ---------- Figure Legends + embedded figures ----------
    figure_legends = {}
    if "Figure Legends" in by_name:
        for kind, payload in parse_paragraph_blocks(by_name["Figure Legends"]):
            if kind == "p":
                num, body = strip_figure_legend(payload)
                if num is not None:
                    figure_legends[num] = body
    for num in range(1, 6):
        legend = figure_legends.get(num, "")
        if not legend:
            print(f"  !! missing legend for Figure {num}")
        img = FIG_FILES[num]
        if os.path.exists(img):
            add_figure_page(doc, f"Figure {num}", img, legend)
        else:
            print(f"  !! missing figure image: {img}")

    # ---------- Supplemental Items (Tables S1-S4 + Fig S1, at the end of the manuscript) ----------
    doc.add_page_break()
    add_heading(doc, "Supplemental Items", 1)
    if "Supplemental Items" in by_name:
        for kind, payload in parse_paragraph_blocks(by_name["Supplemental Items"]):
            if kind == "p":
                if payload.startswith("**Supplemental Figure S1."):
                    legend = payload[len("**Supplemental Figure S1.**"):].strip()
                    add_figure_page(doc, "Supplemental Figure S1", S1_FILE, legend)
                elif payload.startswith("**Table S"):
                    add_caption(doc, payload)
                else:
                    p = doc.add_paragraph(); set_double(p)
                    p.paragraph_format.left_indent = Inches(0.35)
                    add_inline(p, payload)
            elif kind == "table":
                add_three_line_table(doc, payload[0], payload[1:])
            elif kind == "list":
                p = doc.add_paragraph(); set_double(p)
                p.paragraph_format.left_indent = Inches(0.35)
                add_inline(p, payload)

    force_times_new_roman(doc)

    doc.save(OUTP)
    print("Saved:", OUTP)

if __name__ == "__main__":
    main()
